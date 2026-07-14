"""
Generate improved thesis sections 3.2.6 – 3.2.8 as a standalone .docx
that can be copy-pasted / merged into the main chapter document.

Run:
    python3 thesis/generate_improved_sections.py
Output:
    being_worked_on/Chapter3_326_onwards_improved.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy, os

OUT_PATH = "being_worked_on/Chapter3_326_onwards_improved.docx"

# ─── helpers ──────────────────────────────────────────────────────────────────

def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p

def add_code_block(doc, code: str, caption: str = ""):
    """
    Render a code block as a table with one cell, grey background, Courier font.
    Optionally add a caption paragraph below.
    """
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)

    # Grey shading
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    tcPr.append(shd)

    # Code text — split into lines and add each as a separate paragraph
    # so Word doesn't collapse the whitespace.
    lines = code.strip("\n").split("\n")
    first = True
    for line in lines:
        if first:
            para = cell.paragraphs[0]
            first = False
        else:
            para = cell.add_paragraph()
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)

    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.runs[0] if cap.runs else cap.add_run(caption)
        cap_run.font.size = Pt(9)
        cap_run.font.italic = True

    doc.add_paragraph("")   # spacing after block

def add_note(doc, text):
    """Italicised 'Note:' paragraph."""
    p = doc.add_paragraph()
    r = p.add_run("Note: ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
#  Document
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Set default font
style = doc.styles["Normal"]
style.font.name  = "Times New Roman"
style.font.size  = Pt(12)

# ─────────────────────────────────────────────────────────────────────────────
# 3.2.6  Software Architecture
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "3.2.6  Software Architecture", level=2)

body(doc,
    "The gyroscope stabilisation system comprises three software layers that "
    "operate concurrently: an ESP32 firmware application running the real-time "
    "control and sensor loops, a Node.js server that mediates between hardware "
    "and the user interface, and a React dashboard that provides live "
    "telemetry visualisation, session management, and remote parameter tuning. "
    "Each layer is described below together with representative extracts from "
    "the actual production code.")

# ─────────────────────────────────────────────────────────────────────────────
# 3.2.6.1  ESP32 Firmware Architecture
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "3.2.6.1  ESP32 Firmware Architecture (v4.0)", level=3)

body(doc,
    "The firmware (2 004 lines of C++, compiled with the Arduino-ESP32 "
    "framework under PlatformIO) implements a non-blocking cooperative "
    "scheduler inside a single loop() function. Every major subsystem is "
    "polled on each iteration; time-sensitive work is protected by interval "
    "guards rather than blocking delay() calls. The loop executes the "
    "following tasks in order on every iteration:")

steps = [
    ("Wi-Fi / OTA / HTTP polling",
     "ArduinoOTA.handle(), httpServer.handleClient(), and "
     "maintainWifiConnection() keep the over-the-air update path and the "
     "HTTP command server responsive without blocking the control loop."),
    ("Serial and IR remote input",
     "handleSerialInput() drains one byte per iteration from the UART "
     "ring buffer; handleIRRemote() checks the IRremote decode queue. "
     "Both converge on the same executeCommand(char) dispatcher so that "
     "a keyboard keystroke, an IR button press, a Wi-Fi JSON packet, and a "
     "dashboard CMD: string all produce identical results."),
    ("LED keep-alive",
     "A shared ledOffMs timer is extended by every command source "
     "(signalLed()). The LED stays solid for rapid key repeats and goes "
     "low automatically when the timer expires."),
    ("Hardware reset button (debounced)",
     "A 50 ms debounce guard followed by a 500 ms action cool-down "
     "prevents false triggers on GPIO 23."),
    ("IMU read, conversion, and Kalman angle update",
     "readMPU6050() performs an I²C burst read of all 14 registers "
     "(accelerometer X/Y/Z, temperature, gyroscope X/Y/Z) in a single "
     "transaction. convertMPU6050() scales the raw integers to SI units and "
     "subtracts the calibrated gyro bias. kalmanUpdate() then propagates "
     "the two-state filter for each axis."),
    ("Vibration RMS update",
     "updateVibration() appends the current acceleration magnitude to a "
     "64-sample circular buffer and computes the running root-mean-square. "
     "This scalar is reported in the telemetry JSON as vibrationRMS and "
     "is used by the server-side safety monitor. No FFT is performed; the "
     "RMS value captures vibration intensity without requiring spectral "
     "analysis."),
    ("RPM pulse processing",
     "processRPMPulse() reads the interrupt-latched pulse count and timestamp "
     "under a brief interrupt-disable critical section, computes inter-pulse "
     "period, and applies a 75/25 exponential moving average. publishRPMValue() "
     "updates the globally visible rpmValue every RPM_REPORT_INTERVAL_MS."),
    ("Balance control loop (20 ms fixed rate)",
     "If (nowMs − lastCtrlMs) ≥ CTRL_INTERVAL_MS (20 ms) and autoMode && "
     "motorRunning && mpuOk, runBalanceController(dt) is called. The PID "
     "correction is added to manualThrottle and written directly to the ESC "
     "via applyAutoThrottle() — bypassing the ramp to meet the <50 ms "
     "actuator lag requirement imposed by the unstable pole at ~10 rad/s."),
    ("Telemetry emission",
     "sendAppTelemetry() fires every TELEMETRY_INTERVAL_MS (default 250 ms), "
     "serialising ~25 fields as a flat JSON object that is written to both the "
     "115 200-baud UART and the active Wi-Fi TCP socket."),
]

for i, (title, desc) in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(f"{title}. ").bold = True
    p.add_run(desc)

doc.add_paragraph("")

# ── Gyro calibration ──────────────────────────────────────────────────────────
heading(doc, "Gyroscope Bias Calibration", level=4)

body(doc,
    "Before any control loop can run, the gyroscope zero-rate offsets must be "
    "measured so that a stationary sensor reads exactly 0 °/s on all axes. "
    "calibrateGyroBias() collects 200 raw readings at startup, computes the "
    "arithmetic mean, and stores the three offsets in gyroOffsetX/Y/Z. A "
    "second calibration pass is triggered automatically every time the motor "
    "is stopped, since a stationary flywheel provides the best possible "
    "calibration window. Listing 3.1 shows the complete routine.")

add_code_block(doc, """\
void calibrateGyroBias(int samples = 200) {
  Serial.println("[CAL] Collecting gyro bias samples — keep still...");
  double sumX = 0, sumY = 0, sumZ = 0;
  int good = 0;

  for (int i = 0; i < samples; i++) {
    Wire.beginTransmission(MPU6050_ADDR);
    Wire.write(GYRO_XOUT_H);
    if (Wire.endTransmission(false) != 0) { delay(5); continue; }
    Wire.requestFrom(MPU6050_ADDR, 6, 1);
    if (Wire.available() < 6)              { delay(5); continue; }

    int16_t gx = (Wire.read() << 8) | Wire.read();
    int16_t gy = (Wire.read() << 8) | Wire.read();
    int16_t gz = (Wire.read() << 8) | Wire.read();

    // 131 LSB/(°/s) is the sensitivity for the ±250 °/s full-scale range
    sumX += gx / 131.0;
    sumY += gy / 131.0;
    sumZ += gz / 131.0;
    good++;
    delay(2);
  }

  if (good > 0) {
    gyroOffsetX = (float)(sumX / good);
    gyroOffsetY = (float)(sumY / good);
    gyroOffsetZ = (float)(sumZ / good);
    gyroCal = true;
    Serial.printf("[CAL] Offsets: X=%.3f  Y=%.3f  Z=%.3f  (%d samples)\\n",
                  gyroOffsetX, gyroOffsetY, gyroOffsetZ, good);
  }
}""",
"Listing 3.1 — Gyroscope bias calibration (gyro_controller.ino, lines 705–732)")

# ── MPU-6050 initialisation ───────────────────────────────────────────────────
heading(doc, "MPU-6050 Initialisation and DLPF Configuration", level=4)

body(doc,
    "The MPU-6050 is initialised over I²C (SDA = GPIO 21, SCL = GPIO 22) "
    "with three explicit register writes before any measurement is accepted. "
    "Listing 3.2 shows initMPU6050(). The third write — register 0x1A with "
    "value 0x04 — configures the on-chip Digital Low-Pass Filter (DLPF) to "
    "the 20 Hz bandwidth setting. This is a deliberate design decision: the "
    "spinning flywheel generates mechanical vibration in the 25–200 Hz range "
    "that would otherwise corrupt the gyroscope output and cause the Kalman "
    "filter to track vibration instead of rigid-body attitude. The DLPF "
    "attenuates those components before they reach the digital registers. "
    "The accelerometer range is set to ±8 g (register 0x1C, value 0x10) to "
    "avoid saturation during handling, while the gyroscope range is kept at "
    "the highest resolution of ±250 °/s (register 0x1B, value 0x00) because "
    "the system only experiences small angular rates during normal operation.")

add_code_block(doc, """\
void initMPU6050() {
  Wire.beginTransmission(MPU6050_ADDR);
  Wire.write(PWR_MGMT_1);
  Wire.write(0x00);                       // wake from sleep
  if (Wire.endTransmission(true) != 0) {
    Serial.println("[MPU] Init FAILED — not found on I2C bus");
    mpuOk = false; return;
  }
  delay(100);

  Wire.beginTransmission(MPU6050_ADDR);
  Wire.write(ACCEL_CONFIG);
  Wire.write(0x10);                       // ±8 g full-scale
  Wire.endTransmission(true);
  delay(50);

  Wire.beginTransmission(MPU6050_ADDR);
  Wire.write(GYRO_CONFIG);
  Wire.write(0x00);                       // ±250 °/s — highest resolution
  Wire.endTransmission(true);
  delay(50);

  Wire.beginTransmission(MPU6050_ADDR);
  Wire.write(0x1A);
  Wire.write(0x04);                       // DLPF 20 Hz — rejects flywheel vibration
  Wire.endTransmission(true);
  delay(50);

  mpuOk = true;
  Serial.println("[MPU] MPU-6050 OK");
}""",
"Listing 3.2 — MPU-6050 initialisation with DLPF configuration (lines 763–796)")

# ── Kalman filter ─────────────────────────────────────────────────────────────
heading(doc, "Kalman Filter for Tilt Angle Estimation", level=4)

body(doc,
    "Raw accelerometer data gives a noisy but drift-free tilt estimate; the "
    "gyroscope rate gives a smooth but drift-prone integral. The Kalman filter "
    "fuses both signals optimally. The implementation uses a two-state model "
    "with state vector x = [θ, b]ᵀ, where θ is the tilt angle in degrees and "
    "b is the slowly-varying gyro bias. The filter is run independently on the "
    "X and Y axes at every IMU read. Listing 3.3 shows the complete routine "
    "(20 lines of arithmetic, no external library).")

add_code_block(doc, """\
// State: angle (θ) and gyro bias (b).  P[2][2] is the error covariance.
// newAngle  = accelerometer-derived angle  [deg]
// newRate   = raw gyro angular rate        [deg/s]
// dt        = elapsed time since last call [s]

void kalmanUpdate(float &angle, float &bias, float P[2][2],
                  float newAngle, float newRate, float dt) {
  // ── Predict ──────────────────────────────────────────────────────────────
  float rate = newRate - bias;          // bias-corrected angular rate
  angle += dt * rate;                   // integrate rate → predicted angle

  // Propagate covariance  P = F·P·Fᵀ + Q
  P[0][0] += dt * (dt*P[1][1] - P[0][1] - P[1][0] + Q_angle);
  P[0][1] -= dt * P[1][1];
  P[1][0] -= dt * P[1][1];
  P[1][1] += Q_bias * dt;

  // ── Update ───────────────────────────────────────────────────────────────
  float S  = P[0][0] + R_measure;      // innovation variance
  float K0 = P[0][0] / S;              // Kalman gain for angle
  float K1 = P[1][0] / S;              // Kalman gain for bias
  float y  = newAngle - angle;         // innovation (measurement residual)

  angle += K0 * y;                     // correct angle estimate
  bias  += K1 * y;                     // correct bias estimate

  // Update covariance  P = (I − K·H)·P
  float P00=P[0][0], P01=P[0][1], P10=P[1][0], P11=P[1][1];
  P[0][0] = P00 - K0*P00;
  P[0][1] = P01 - K0*P01;
  P[1][0] = P10 - K1*P00;
  P[1][1] = P11 - K1*P01;
}""",
"Listing 3.3 — Two-state Kalman filter for tilt estimation (lines 738–757)")

body(doc,
    "The filter noise parameters Q_angle = 0.001, Q_bias = 0.003, and "
    "R_measure = 0.03 were tuned empirically. A larger R_measure makes the "
    "filter trust the gyroscope more heavily (smoother but more prone to "
    "drift); a smaller value gives the accelerometer more weight (less drift "
    "but more vibration noise). The chosen values give the filter an effective "
    "crossover frequency of approximately 5 Hz, well above the rigid-body "
    "tilt dynamics and well below the flywheel vibration band suppressed by "
    "the DLPF.")

# ── PID balance controller ────────────────────────────────────────────────────
heading(doc, "v4 Balance Controller", level=4)

body(doc,
    "The central contribution of firmware v4 is a redesigned PID controller "
    "that addresses three documented failures in v3.5. The file header records "
    "the reasoning explicitly:")

failures = [
    ("Insufficient gains (v3.5)",
     "v3.5 used Kp = 0.7 and Ki = 0.005, both below the Routh stability "
     "boundary derived from the reaction-wheel inverted-pendulum plant model "
     "(Ki_min ≈ 17, Kp_min ≈ 1.4). The controller was mathematically unstable "
     "by design. v4 defaults are Kp = 3.0, Ki = 25.0, Kd = 1.0."),
    ("Noisy derivative term (v3.5)",
     "v3.5 computed D = Kd·(error − prevError)/dt, which numerically "
     "differentiates the Kalman angle estimate and amplifies quantisation noise. "
     "v4 uses the raw gyroscope angular rate directly: D = −Kd·ω_gyro, "
     "eliminating differentiation entirely."),
    ("Single-axis correction only (v3.5)",
     "v3.5 only corrected whichever of X or Y had the larger error at each "
     "tick. v4 maintains separate integral accumulators for both axes and "
     "blends them into a single throttle correction."),
]

for i, (title, desc) in enumerate(failures, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(f"{title}. ").bold = True
    p.add_run(desc)

doc.add_paragraph("")

body(doc,
    "The control law implemented by runBalanceController() is reproduced "
    "verbatim in the firmware file header comment (lines 1184–1201) and is "
    "given here for reference:")

add_code_block(doc, """\
Control law (per axis i ∈ {X, Y}):

  e_i    =  setpoint_i − tilt_i                    [deg]
  I_i   +=  Ki · e_i · dt     (frozen when saturated)
  P_i    =  Kp · e_i
  D_i    =  −Kd · ω_gyro_i                         [direct gyro rate]
  NL_i   =  Knl · sign(e_i) · √(|e_i| − NL_THRESH)
             if |e_i| > NL_THRESH, else 0
  u_i    =  P_i + I_i + D_i + NL_i

Combined output:
  dominant axis = argmax(|e_X|, |e_Y|)
  if both |e| > 2°:  u = u_dom + u_sub · (|e_sub|/|e_dom|) · 0.4
  else:              u = u_dom

ESC pulse = manualThrottle + u   (clamped to [MIN_THROTTLE, MAX_THROTTLE])""",
"Figure 3.x — Control law summary (firmware header, lines 1184–1201)")

body(doc,
    "Listing 3.4 shows the C++ implementation. The anti-windup logic is "
    "particularly noteworthy: integration is frozen not only when the ESC "
    "rail is saturated (outputSaturated flag) but also when the pre-saturation "
    "candidate correction is within 5 % of the MAX_CORRECTION_US clamp. This "
    "back-calculation scheme prevents slow integrator wind-up in the regime "
    "just short of hard saturation, which is the most common operating "
    "condition during a large disturbance recovery.")

add_code_block(doc, """\
static float runBalanceController(float dt) {
  if (!mpuOk) return 0.0f;

  float eX = setpointX - error_X;    // tilt error per axis [deg]
  float eY = setpointY - error_Y;
  float absX = fabsf(eX), absY = fabsf(eY);

  // ── Proportional ──────────────────────────────────────────────────────────
  float pX = bal_Kp * eX;
  float pY = bal_Kp * eY;

  // ── Integral with anti-windup ─────────────────────────────────────────────
  bool correctionClamped =
    outputSaturated ||
    (fabsf(integralX + bal_Kp*eX - bal_Kd*gyroX_dps) >= MAX_CORRECTION_US*0.95f) ||
    (fabsf(integralY + bal_Kp*eY - bal_Kd*gyroY_dps) >= MAX_CORRECTION_US*0.95f);
  if (!correctionClamped) {
    integralX += bal_Ki * eX * dt;
    integralY += bal_Ki * eY * dt;
  }
  // Hard clamp: integrals never exceed ½ saturation limit
  integralX = constrain(integralX, -MAX_CORRECTION_US*0.5f, MAX_CORRECTION_US*0.5f);
  integralY = constrain(integralY, -MAX_CORRECTION_US*0.5f, MAX_CORRECTION_US*0.5f);

  // ── Derivative: direct gyro rate (no numerical differentiation) ───────────
  float dX = -bal_Kd * gyroX_dps;
  float dY = -bal_Kd * gyroY_dps;

  // ── Non-linear boost for large disturbances ───────────────────────────────
  float nlX = 0.0f, nlY = 0.0f;
  if (absX > NL_THRESH_DEG)
    nlX = bal_Knl * (eX > 0.0f ? 1.0f : -1.0f) * sqrtf(absX - NL_THRESH_DEG);
  if (absY > NL_THRESH_DEG)
    nlY = bal_Knl * (eY > 0.0f ? 1.0f : -1.0f) * sqrtf(absY - NL_THRESH_DEG);

  // ── Per-axis totals ───────────────────────────────────────────────────────
  float uX = pX + integralX + dX + nlX;
  float uY = pY + integralY + dY + nlY;

  // ── Two-axis blending ─────────────────────────────────────────────────────
  float correction;
  const float BOTH_AXIS_THRESHOLD = 2.0f;
  if (absX >= absY)
    correction = (absY > BOTH_AXIS_THRESHOLD) ? uX + uY*(absY/absX)*0.4f : uX;
  else
    correction = (absX > BOTH_AXIS_THRESHOLD) ? uY + uX*(absX/absY)*0.4f : uY;

  // ── Clamp output and update saturation flag ───────────────────────────────
  correction = constrain(correction, -MAX_CORRECTION_US, MAX_CORRECTION_US);
  float candidate = (float)manualThrottle + correction;
  outputSaturated = (candidate <= (float)MIN_THROTTLE ||
                     candidate >= (float)MAX_THROTTLE);

  return correction;
}""",
"Listing 3.4 — v4 balance controller with anti-windup and two-axis blending (lines 1204–1281)")

# ── RPM ISR ───────────────────────────────────────────────────────────────────
heading(doc, "RPM Sensing via Hardware Interrupt", level=4)

body(doc,
    "Flywheel speed is measured by a Hall-effect sensor on GPIO 16. The "
    "sensor produces one falling edge per revolution. The interrupt service "
    "routine (ISR) is placed in IRAM with the IRAM_ATTR attribute so it "
    "executes from RAM rather than flash, guaranteeing deterministic latency "
    "even during cache misses caused by Wi-Fi DMA activity.")

add_code_block(doc, """\
// IRAM_ATTR: stored in internal RAM — executes even during flash cache misses
void IRAM_ATTR rpmPulseISR() {
  rpmPulseCount++;                        // atomic on Xtensa (32-bit write)
  rpmPulseTimestampUs = micros();         // record time of this pulse
}

static void processRPMPulse() {
  // Snapshot interrupt state atomically (disable interrupts for two reads)
  uint32_t snapshotCount = 0, snapshotTimeUs = 0;
  noInterrupts();
  snapshotCount  = rpmPulseCount;
  snapshotTimeUs = rpmPulseTimestampUs;
  interrupts();

  if (snapshotCount != lastRpmPulseCount && snapshotTimeUs > 0) {
    uint32_t deltaUs = snapshotTimeUs - lastRpmPulseMicros;
    if (deltaUs > 0 && deltaUs < 6000000UL) {         // sanity: < 10 RPM floor
      float cand = 60000000.0f / (float)(deltaUs * RPM_PULSES_PER_REV);
      if (cand >= RPM_MIN_VALID_RPM && cand <= RPM_MAX_VALID_RPM)
        // 75/25 exponential moving average for smoothing
        currentRPM = currentRPM > 0.0f ? (currentRPM*0.75f + cand*0.25f) : cand;
    }
    lastRpmPulseMicros = snapshotTimeUs;
    lastRpmPulseCount  = snapshotCount;
  }

  // RPM timeout — declare zero if no pulse within RPM_TIMEOUT_MS
  if ((micros() - lastRpmActivityMicros) > RPM_TIMEOUT_MS * 1000UL)
    currentRPM = 0.0f;
}""",
"Listing 3.5 — RPM interrupt service routine and pulse processing (lines 1059–1093)")

add_note(doc,
    "The control loop does not gate on RPM. Auto mode applies PID correction "
    "on every tick regardless of whether a valid RPM reading is available. "
    "RPM is telemetry only; it does not pause or cancel stabilisation. This is "
    "a deliberate choice because the RPM sensor updates at 250 ms cadence "
    "(one pulse per revolution at 240 RPM), which is too slow to be a "
    "stabilisation signal for a 20 ms control loop.")

doc.add_paragraph("")

# ── Telemetry JSON ────────────────────────────────────────────────────────────
heading(doc, "Telemetry JSON Format", level=4)

body(doc,
    "Every TELEMETRY_INTERVAL_MS milliseconds (default 250 ms) the firmware "
    "serialises ~25 fields as a single flat JSON object and writes it "
    "simultaneously to the USB UART and the active Wi-Fi TCP socket (port 5001). "
    "A representative payload is shown below. The server ingests this via two "
    "independent code paths (USB serial bridge and Wi-Fi TCP server) that both "
    "call the same ingestReading() function.")

add_code_block(doc, """\
{
  "mode":          "auto",        // "auto" | "manual"
  "tiltX":         -1.23,         // Kalman tilt error from reference, X axis [deg]
  "tiltY":          0.45,         // Kalman tilt error from reference, Y axis [deg]
  "filteredAngleX":-1.23,         // mirrors tiltX (post-reference; raw angle before zeroing)
  "filteredAngleY": 0.45,
  "tiltZ":          0.00,         // Z-axis (yaw) placeholder
  "rpm":           2847,          // filtered flywheel speed [RPM]
  "temp":          31.50,         // MPU-6050 die temperature fallback [°C]
  "pwm":           42.30,         // ESC duty as percentage of usable range
  "throttle":      1341,          // current ESC pulse width [µs]
  "escPulse":      1341.0,        // same as throttle (dedicated field)
  "heldFINALpulse":   0.5,        // correction term scaled to µs/10
  "SentMOTORpulse":1341.0,        // last value actually written to LEDC
  "vibrationRMS":  0.0032,        // RMS of 64-sample accel magnitude buffer
  "correction":   -12.3,          // raw PID output before ESC clamping [arb]
  "intX":          -3.21,         // integral accumulator, X axis
  "intY":           1.05,         // integral accumulator, Y axis
  "saturated":     false,         // true when ESC output rail is saturated
  "kp":             7.0,          // live gains (updated by dashboard sliders)
  "ki":             1.0,
  "kd":             1.5,
  "knl":            1.2,
  "mpuFault":      false          // true if I2C burst read failed
}""",
"Listing 3.6 — Representative firmware telemetry payload (sendAppTelemetry(), lines 1771–1824)")

body(doc,
    "The Wi-Fi path uses the same newline-delimited text protocol as USB serial. "
    "If the TCP socket is connected (tcpClient.connected()), the JSON line is "
    "written to both Serial and the socket in a single sendAppTelemetry() call, "
    "eliminating any skew between the two streams.")

# ─────────────────────────────────────────────────────────────────────────────
# 3.2.6.2  Node.js Server Architecture
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "3.2.6.2  Node.js Server Architecture", level=3)

body(doc,
    "The Node.js server (Express + WebSocket, running on port 8080) acts as "
    "the single integration point between the physical hardware and the web "
    "dashboard. It performs six concurrent tasks:")

server_tasks = [
    ("Dual-path hardware ingestion",
     "Sensor data arrives over one of two independent channels. The USB "
     "channel is managed by a companion Electron desktop application "
     "(the 'serial bridge') that reads the UART, parses each line, and posts "
     "readings to POST /api/readings/ingest. The Wi-Fi channel is handled "
     "directly by the server: a raw TCP server (wifi-tcp-server.ts) listens "
     "on port 5001, parses the same newline-delimited protocol, and calls the "
     "same ingestReading() function internally. Both paths produce identical "
     "database writes and WebSocket broadcasts."),
    ("Device capability negotiation",
     "On connection (USB or Wi-Fi), the firmware sends a DEVICE_INFO handshake "
     "line: DEVICE_INFO:esp32WiFi|VERSION:4.0|COMPONENTS:IR,MPU6050,ESC,RPM_SENSOR|BAUD:115200. "
     "The server's componentsToCaps() function maps the comma-separated "
     "component list to a typed capability object that controls which "
     "dashboard panels are shown or hidden (e.g., if the RPM sensor flag is "
     "absent, the RPM gauge is replaced by a placeholder)."),
    ("WebSocket broadcast to dashboard clients",
     "Every ingested reading, alert, session update, and motor-state change is "
     "pushed to all connected browser clients via broadcastToClients(). On "
     "initial WebSocket connection the server immediately emits a "
     "hardware_connected event carrying the current hardware status, so the "
     "dashboard does not have to wait up to 250 ms for the next telemetry "
     "frame before knowing the device is live."),
    ("Safety monitoring",
     "Each ingest call evaluates the incoming reading against configurable "
     "thresholds stored in the settings store (tempWarnThreshold, "
     "tempCritThreshold, vibWarnThreshold, vibCritThreshold). A temperature "
     "or vibration breach at the critical level triggers an immediate emergency "
     "stop: the server writes CMD:EMERGENCY_STOP to the hardware queue, "
     "updates motor state, fires a critical-level alert, and closes the running "
     "session record."),
    ("Wi-Fi connection lifecycle management",
     "The TCP server implements a three-state socket lifecycle. A fresh "
     "connection is placed in a 'pending' state for up to 60 seconds; the "
     "dashboard shows an Accept / Reject banner. On accept, the socket is "
     "activated. If the link drops mid-session, an 8-second grace window is "
     "opened — a reconnect within that window resumes the session without "
     "interruption; if no reconnect arrives the session is closed and a "
     "hardware_disconnected event is broadcast."),
    ("REST API for dashboard queries",
     "Forty-plus endpoints (defined in the OpenAPI spec at lib/api-spec/openapi.yaml) "
     "serve session history, alert lists, SD-card data, OTA firmware uploads, "
     "settings persistence, and IR command logs. The dashboard consumes all of "
     "these through auto-generated TypeScript query hooks derived from the spec "
     "at build time."),
]

for i, (title, desc) in enumerate(server_tasks, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(f"{title}. ").bold = True
    p.add_run(desc)

doc.add_paragraph("")

body(doc,
    "Listing 3.7 shows the Wi-Fi TCP server's socket data handler, which "
    "illustrates the line-parsing architecture that is shared (conceptually) "
    "by both ingestion paths.")

add_code_block(doc, """\
// wifi-tcp-server.ts — socket data handler (lines 424–429)
socket.on("data", (chunk: string) => {
  buf += chunk;
  const lines = buf.split("\\n");
  buf = lines.pop() ?? "";          // keep any partial (unterminated) line
  lines.forEach(line => parseLine(line, clientIp));
});

// parseLine() dispatches on the first token:
//   "{"           → JSON.parse → ingestReading()   (telemetry)
//   "DEVICE_INFO" → capability handshake
//   "STATUS:"     → forward to dashboard as info alert
//   "CMD:IR_"     → forward to IR command feed
function parseLine(line: string, clientIp: string): void {
  line = line.trim();
  if (!line) return;
  if (line.startsWith("{")) {
    const data = JSON.parse(line) as Record<string, unknown>;
    if (data["status"] === "boot") return;   // boot message, not a reading
    void ingestReading(data);
    return;
  }
  if (line.startsWith("DEVICE_INFO:")) { /* ... parse handshake ... */ return; }
  if (line.startsWith("STATUS:"))      { broadcastToClients({ type: "alert", ... }); return; }
  if (line.startsWith("CMD:IR_"))      { broadcastToClients({ type: "ir_command", ... }); return; }
}""",
"Listing 3.7 — Wi-Fi TCP socket data handler and line parser (wifi-tcp-server.ts, lines 420–376)")

body(doc,
    "Command delivery follows the reverse path. When the dashboard sends a "
    "motor command, the server first attempts sendWifiCommand() (direct TCP "
    "write); if the Wi-Fi socket is absent or torn down it falls back to "
    "queueSerialWrite(), which the Electron bridge drains via "
    "GET /api/serial/pending-write on a polling interval. This dual-path "
    "command delivery ensures commands reach the hardware regardless of which "
    "connectivity mode is active.")

# ─────────────────────────────────────────────────────────────────────────────
# 3.2.6.3  React Dashboard Architecture
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "3.2.6.3  React Dashboard Architecture", level=3)

body(doc,
    "The dashboard is a React 18 single-page application built with Vite. "
    "It communicates with the server through two complementary channels: "
    "a WebSocket connection that delivers push events (readings, alerts, "
    "session updates, hardware state changes), and a REST API consumed through "
    "auto-generated TypeScript query hooks. The OpenAPI specification "
    "(lib/api-spec/openapi.yaml, ~1 050 lines) is the single source of truth "
    "for all API contracts; the client SDK is regenerated at build time via "
    "openapi-typescript-codegen, ensuring the dashboard and server never diverge "
    "on field names or types.")

body(doc,
    "The dashboard is organised into five functional sections:")

dash_sections = [
    ("Live monitoring panel",
     "Three real-time charts (tilt X/Y, RPM, ESC pulse width) update at the "
     "WebSocket read rate (~250 ms). A tilt bubble shows the 2D orientation "
     "vector. A session timer and motor state badge provide at-a-glance status."),
    ("Motor control and PID tuning",
     "Start/Stop/Emergency-Stop buttons dispatch motor commands. A PID slider "
     "panel sends CMD:SET_PID:<kp>,<ki>,<kd> to the firmware in real time, "
     "allowing gain tuning without reflashing. The same tuning effect can be "
     "achieved from the IR remote (buttons 1–6) or the serial keyboard (p/o/i/u/d/c)."),
    ("Alert centre",
     "A five-level alert system (info / warning / error / critical / emergency) "
     "displays server-generated safety events. Alerts are delivered over "
     "WebSocket and also polled from the database. User preferences (which "
     "levels produce a notification sound, which auto-dismiss) are persisted in "
     "localStorage. Each alert level is independently acknowledgeable."),
    ("Session management",
     "A session dialog (triggered by the dashboard Start button or IR CH+ "
     "key) creates a database record that groups all readings and alerts "
     "within a named run. Session history is searchable and exportable as CSV."),
    ("Firmware and connectivity management",
     "A firmware panel accepts OTA binary uploads (forwarded via the server to "
     "the ESP32's ArduinoOTA handler). A connectivity dialog shows the "
     "active channel (USB serial port or Wi-Fi TCP), firmware version, "
     "and component health flags parsed from the DEVICE_INFO / HEALTH lines. "
     "A pending-connection banner surfaces when an ESP32 attempts a first "
     "Wi-Fi TCP connection, allowing the user to Accept or Reject it."),
]

for i, (title, desc) in enumerate(dash_sections, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(f"{title}. ").bold = True
    p.add_run(desc)

doc.add_paragraph("")
body(doc, "[Figure 3.x — Dashboard screenshot showing live monitoring panel — insert screenshot here]")
doc.add_paragraph("")

# ─────────────────────────────────────────────────────────────────────────────
# 3.2.7  MATLAB Simulink Simulation
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "3.2.7  MATLAB Simulink Simulation", level=2)

body(doc,
    "Three Simulink simulation scenarios were constructed to evaluate the "
    "control design before hardware implementation: a step-response test, "
    "a disturbance rejection test, and a parameter sensitivity sweep.")

body(doc,
    "[Figure 3.x — Simulink block diagram — insert figure here]")
body(doc,
    "[Figure 3.x — Step response simulation result — insert figure here]")
body(doc,
    "[Figure 3.x — Disturbance rejection result — insert figure here]")

doc.add_paragraph("")

# ── IMPORTANT: gain discrepancy note ─────────────────────────────────────────
heading(doc, "Relationship between Simulation Gains and Deployed Gains", level=3)

body(doc,
    "An important distinction must be drawn between the PID gains used in the "
    "Simulink model and those deployed in the v4 firmware. The Simulink plant "
    "was an RPM-control model: the controlled variable was flywheel speed and "
    "the PID output commanded a throttle increment to reach a target RPM. The "
    "Simulink-optimised gains (Kp = 0.85, Ki = 2.40, Kd = 0.015) are therefore "
    "expressed in units appropriate for that plant (RPM error → throttle µs).")

body(doc,
    "Firmware v4, by contrast, controls tilt angle, not RPM. The controlled "
    "variable is the Kalman-filtered tilt error in degrees; the PID output is a "
    "correction added to the base ESC pulse in microseconds. This is a "
    "fundamentally different plant with different gain units, different loop "
    "dynamics, and different Routh stability boundaries. The Routh criterion "
    "for the reaction-wheel inverted-pendulum plant (mass 0.5 kg, effective "
    "moment arm L, motor constant K_eff) requires:")

add_code_block(doc, """\
Ki_min  ≈  m · g · L / K_eff  ≈  17 – 20   [deg·s / µs]
Kp_min  ≈  m · g · L · τ / K_eff  ≈  1.4    [µs / deg]

v3.5 deployed gains:  Kp = 0.7,   Ki = 0.005   ← below both Routh minimums
v4   deployed gains:  Kp = 7.0,   Ki = 1.0,  Kd = 1.5, Knl = 1.2""",
"Table — Gain comparison: Simulink (RPM plant) vs. deployed firmware (tilt plant)")

body(doc,
    "The Simulink work therefore served two purposes: it validated the "
    "conceptual feasibility of gyroscopic stabilisation using reaction-wheel "
    "torque, and it provided a systematic starting point for understanding "
    "controller structure (proportional, integral, and derivative terms; "
    "anti-windup). The numerical gain values were not transferred to the "
    "hardware controller directly, because the plant model changed. The v4 "
    "gains were determined by applying the Routh stability criterion to the "
    "tilt-angle plant and then refined through hardware iteration, as "
    "documented in Appendix [X].")

# ─────────────────────────────────────────────────────────────────────────────
# 3.2.8  Circuit Design and Simulation
# ─────────────────────────────────────────────────────────────────────────────
heading(doc, "3.2.8  Circuit Design and Simulation", level=2)

body(doc,
    "The circuit was designed and simulated in Proteus. Five sub-circuits "
    "correspond to the five main hardware blocks: the ESP32 microcontroller, "
    "the MPU-6050 IMU, the ESC/motor drive path, the RPM sensor, and the "
    "IR remote receiver. Each is discussed below together with the firmware "
    "register configuration that validates the design choice.")

body(doc, "[Figure 3.x — Proteus top-level schematic — insert figure here]")
doc.add_paragraph("")

# ── ESC / LEDC ────────────────────────────────────────────────────────────────
heading(doc, "ESC Signal Generation via ESP32 LEDC", level=3)

body(doc,
    "The electronic speed controller expects a standard RC PWM signal: "
    "50 Hz carrier, pulse width between 1 000 µs (minimum throttle) and "
    "2 000 µs (maximum throttle). The ESP32 generates this signal using its "
    "LED Control (LEDC) peripheral, which provides a 16-bit hardware PWM "
    "generator with a dedicated timer independent of the CPU. The firmware "
    "initialises the LEDC in setup() as shown in Listing 3.8. The 16-bit "
    "resolution at 50 Hz gives 65 535 counts per 20 ms period, corresponding "
    "to a pulse-width resolution of approximately 0.3 µs per count — "
    "well below the 1 µs granularity that ESCs can distinguish.")

add_code_block(doc, """\
// setup() — ESC PWM initialisation (lines 1879–1896)
#define ESC_PIN        4          // GPIO 4 → ESC signal wire
#define PWM_CHANNEL    0          // LEDC channel 0
#define PWM_FREQ       50         // 50 Hz carrier (standard RC servo/ESC)
#define PWM_RESOLUTION 16         // 16-bit: 65535 counts per 20 ms period

ledcSetup(PWM_CHANNEL, PWM_FREQ, PWM_RESOLUTION);   // configure timer
ledcAttachPin(ESC_PIN, PWM_CHANNEL);                 // bind GPIO to timer
setThrottle(ARM_THROTTLE);                           // send arming pulse (1000 µs)

// ── ESC arming sequence (3 s at ARM_THROTTLE) ──────────────────────────────
unsigned long escArmStart = millis();
while (millis() - escArmStart < 3000) {
  yield();                  // keep Wi-Fi/OTA/HTTP alive during arming
  handleSerialInput();
  handleIRRemote();
  ArduinoOTA.handle();
  httpServer.handleClient();
}
setThrottle(ARM_THROTTLE);
Serial.println("[ESC] Armed. STATUS:MOTOR=OFF");

// ── Pulse-width to LEDC count conversion ───────────────────────────────────
static void setThrottle(int us) {
  // Period = 1/50 Hz = 20 000 µs  → duty = (us / 20000) × 65535
  uint32_t duty = (uint32_t)((float)us / 20000.0f * 65535.0f);
  ledcWrite(PWM_CHANNEL, duty);
  motorThrottle = us;
}""",
"Listing 3.8 — LEDC PWM setup and pulse-width conversion (setup(), lines 1879–1896)")

body(doc,
    "During normal stop (stopMotor()), the firmware ramps the throttle back to "
    "ARM_THROTTLE smoothly over multiple loop iterations. During an emergency "
    "stop (emergencyStop()), setThrottle(ARM_THROTTLE) is called immediately "
    "and unconditionally — the ramp is bypassed to ensure the fastest possible "
    "ESC response. This distinction is visible in the circuit-level signal: an "
    "oscilloscope probing the ESC signal pin shows a gradual pulse-width "
    "decrease on a normal stop and an immediate step to 1 000 µs on an "
    "emergency stop.")

# ── MPU-6050 I2C ──────────────────────────────────────────────────────────────
heading(doc, "MPU-6050 I²C Connection and Register Map", level=3)

body(doc,
    "The MPU-6050 connects to the ESP32 via I²C on the standard Arduino-ESP32 "
    "pins (SDA = GPIO 21, SCL = GPIO 22). The device address is 0x68 (AD0 pin "
    "pulled low). The firmware performs a 14-byte burst read starting at "
    "register 0x3B (ACCEL_XOUT_H) to retrieve all six sensor axes and the "
    "temperature in a single I²C transaction, minimising bus overhead at the "
    "IMU poll rate. Table 3.x maps the key register writes made during "
    "initialisation to their physical effects.")

add_code_block(doc, """\
Register  Value  Effect
────────  ─────  ──────────────────────────────────────────────────────────────
0x6B      0x00   PWR_MGMT_1: wake the device from sleep
0x1C      0x10   ACCEL_CONFIG: ±8 g full-scale → 4096 LSB/g sensitivity
0x1B      0x00   GYRO_CONFIG:  ±250 °/s full-scale → 131 LSB/(°/s) sensitivity
0x1A      0x04   CONFIG (DLPF_CFG=4): 20 Hz low-pass filter bandwidth

Burst-read layout (14 bytes starting at 0x3B):
  Bytes 0–1   ACCEL_XOUT   (signed 16-bit, big-endian)
  Bytes 2–3   ACCEL_YOUT
  Bytes 4–5   ACCEL_ZOUT
  Bytes 6–7   TEMP_OUT     → T[°C] = raw/340 + 36.53
  Bytes 8–9   GYRO_XOUT    (signed 16-bit, big-endian)
  Bytes 10–11 GYRO_YOUT
  Bytes 12–13 GYRO_ZOUT""",
"Table 3.x — MPU-6050 register configuration and burst-read layout")

body(doc,
    "The choice of ±8 g rather than the minimum ±2 g range is justified by "
    "the handling conditions: the assembly is hand-carried and placed on a "
    "bench, where incidental accelerations can easily exceed ±2 g. Using ±8 g "
    "prevents accelerometer saturation at the cost of reduced sensitivity "
    "(4 096 instead of 16 384 LSB/g), which is acceptable since the Kalman "
    "filter primarily uses the accelerometer for static tilt angle "
    "(atan2 from gravity vector) rather than for dynamic acceleration tracking. "
    "The gyroscope remains at ±250 °/s because the rigid-body tilt rates "
    "seen during stabilisation are well within ±10 °/s, and the higher "
    "sensitivity (131 LSB/(°/s)) reduces quantisation error in the Kalman "
    "predict step.")

# ── RPM sensor ────────────────────────────────────────────────────────────────
heading(doc, "RPM Sensor Interface", level=3)

body(doc,
    "The Hall-effect RPM sensor is wired to GPIO 16 with an internal pull-up "
    "resistor enabled by the firmware (pinMode(RPM_SENSOR_PIN, INPUT_PULLUP)). "
    "One falling edge is generated per revolution. The interrupt is attached "
    "with FALLING edge trigger:")

add_code_block(doc, """\
attachInterrupt(digitalPinToInterrupt(RPM_SENSOR_PIN), rpmPulseISR, FALLING);""",
"Listing 3.9 — RPM interrupt attachment (setup(), line 1856)")

body(doc,
    "The Proteus simulation of the RPM sensor path uses a voltage-controlled "
    "square-wave generator to model pulses at representative speeds. "
    "At 3 000 RPM the pulse period is 20 ms, giving one interrupt every "
    "20 000 µs. The inter-pulse timestamp arithmetic in processRPMPulse() "
    "rejects any candidate RPM below RPM_MIN_VALID_RPM or above "
    "RPM_MAX_VALID_RPM as a noise guard. Periods longer than 6 000 000 µs "
    "(10 RPM equivalent) are also rejected, preventing a stale micros() "
    "timestamp from producing a spurious reading after the flywheel stops.")

body(doc, "[Figure 3.x — Proteus RPM sensor simulation circuit — insert figure here]")
body(doc, "[Figure 3.x — Proteus ESC PWM output waveform — insert figure here]")
body(doc, "[Figure 3.x — Proteus MPU-6050 I2C transaction capture — insert figure here]")

# ── save ──────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f"Saved → {OUT_PATH}")
